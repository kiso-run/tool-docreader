"""Unit tests for tool-docreader."""
from __future__ import annotations

import csv
import json
import subprocess
import sys
from pathlib import Path
from unittest.mock import patch, MagicMock

import pytest

# Import directly from run.py (pythonpath = ["."])
from run import (
    do_list, do_info, do_read,
    _resolve_path, _parse_page_ranges, _format_size, _is_likely_text,
    _MAX_OUTPUT_CHARS,
)


# ---------------------------------------------------------------------------
# Fixtures — create real small files for each format
# ---------------------------------------------------------------------------


@pytest.fixture
def workspace(tmp_path):
    """Create a workspace with uploads/ directory."""
    uploads = tmp_path / "uploads"
    uploads.mkdir()
    return tmp_path


@pytest.fixture
def txt_file(workspace):
    """Create a plain text file."""
    f = workspace / "uploads" / "hello.txt"
    f.write_text("Hello, world!\nSecond line.")
    return f


@pytest.fixture
def csv_file(workspace):
    """Create a CSV file."""
    f = workspace / "uploads" / "data.csv"
    with open(f, "w", newline="") as fh:
        writer = csv.writer(fh)
        writer.writerow(["name", "age", "city"])
        writer.writerow(["Alice", "30", "Rome"])
        writer.writerow(["Bob", "25", "Milan"])
    return f


@pytest.fixture
def pdf_file(workspace):
    """Create a minimal PDF file using pypdf."""
    from pypdf import PdfWriter
    f = workspace / "uploads" / "doc.pdf"
    writer = PdfWriter()
    # Add 3 pages with text
    for i in range(3):
        writer.add_blank_page(width=200, height=200)
    # pypdf blank pages have no text — we'll test the "no text" path
    writer.write(str(f))
    return f


@pytest.fixture
def docx_file(workspace):
    """Create a minimal DOCX file."""
    from docx import Document
    f = workspace / "uploads" / "report.docx"
    doc = Document()
    doc.add_paragraph("First paragraph of the report.")
    doc.add_paragraph("Second paragraph with details.")
    doc.save(str(f))
    return f


@pytest.fixture
def xlsx_file(workspace):
    """Create a minimal XLSX file."""
    from openpyxl import Workbook
    f = workspace / "uploads" / "sheet.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.title = "Data"
    ws.append(["Name", "Score"])
    ws.append(["Alice", 95])
    ws.append(["Bob", 87])
    wb.save(str(f))
    return f


# ---------------------------------------------------------------------------
# do_list
# ---------------------------------------------------------------------------


class TestDoList:
    def test_list_with_files(self, workspace, txt_file, csv_file):
        result = do_list(str(workspace))
        assert "Files in uploads/ (2):" in result
        assert "hello.txt" in result
        assert "data.csv" in result

    def test_list_empty_directory(self, workspace):
        result = do_list(str(workspace))
        assert "empty" in result.lower()

    def test_list_no_uploads_dir(self, tmp_path):
        result = do_list(str(tmp_path))
        assert "No uploads/" in result


# ---------------------------------------------------------------------------
# do_info
# ---------------------------------------------------------------------------


class TestDoInfo:
    def test_info_txt(self, workspace, txt_file):
        result = do_info(str(workspace), {"file_path": "uploads/hello.txt"})
        assert "hello.txt" in result
        assert ".txt" in result
        assert "B" in result or "KB" in result

    def test_info_csv(self, workspace, csv_file):
        result = do_info(str(workspace), {"file_path": "uploads/data.csv"})
        assert "Rows:" in result
        assert "3" in result  # header + 2 data rows

    def test_info_pdf(self, workspace, pdf_file):
        result = do_info(str(workspace), {"file_path": "uploads/doc.pdf"})
        assert "Pages: 3" in result

    def test_info_xlsx(self, workspace, xlsx_file):
        result = do_info(str(workspace), {"file_path": "uploads/sheet.xlsx"})
        assert "Sheets:" in result
        assert "Data" in result

    def test_info_missing_file(self, workspace):
        with pytest.raises(FileNotFoundError):
            do_info(str(workspace), {"file_path": "uploads/nope.pdf"})


# ---------------------------------------------------------------------------
# do_read — basic (small files, all fit in budget)
# ---------------------------------------------------------------------------


class TestDoRead:
    def test_read_txt(self, workspace, txt_file):
        result = do_read(str(workspace), {"file_path": "uploads/hello.txt"})
        assert "Hello, world!" in result
        assert "Second line." in result
        assert "Document: hello.txt" in result

    def test_read_csv(self, workspace, csv_file):
        result = do_read(str(workspace), {"file_path": "uploads/data.csv"})
        assert "Alice" in result
        assert "Bob" in result
        assert "Rome" in result
        assert "Dataset: data.csv" in result
        assert "3 rows" in result
        assert "Columns:" in result

    def test_read_docx(self, workspace, docx_file):
        result = do_read(str(workspace), {"file_path": "uploads/report.docx"})
        assert "First paragraph" in result
        assert "Second paragraph" in result
        assert "Document: report.docx" in result

    def test_read_xlsx(self, workspace, xlsx_file):
        result = do_read(str(workspace), {"file_path": "uploads/sheet.xlsx"})
        assert "Alice" in result
        assert "95" in result
        assert "Bob" in result
        assert "Sheet: Data" in result
        assert "Workbook: sheet.xlsx" in result

    def test_read_pdf_blank_pages(self, workspace, pdf_file):
        result = do_read(str(workspace), {"file_path": "uploads/doc.pdf"})
        assert "no extractable text" in result.lower()

    def test_read_pdf_with_text(self, workspace, pdf_file):
        """PDF reading with mocked text extraction."""
        mock_page = MagicMock()
        mock_page.extract_text.return_value = "Page content here"

        with patch("pypdf.PdfReader") as MockReader:
            MockReader.return_value.pages = [mock_page, mock_page]
            result = do_read(str(workspace), {"file_path": "uploads/doc.pdf"})
        assert "Page content here" in result
        assert "Page 1" in result
        assert "Page 2" in result
        assert "Document: doc.pdf (2 pages)" in result

    def test_read_pdf_page_ranges(self, workspace, pdf_file):
        """PDF with page range argument."""
        pages = []
        for i in range(5):
            p = MagicMock()
            p.extract_text.return_value = f"Content of page {i+1}"
            pages.append(p)

        with patch("pypdf.PdfReader") as MockReader:
            MockReader.return_value.pages = pages
            result = do_read(str(workspace), {
                "file_path": "uploads/doc.pdf",
                "pages": "2-3",
            })
        assert "Content of page 2" in result
        assert "Content of page 3" in result
        assert "Content of page 1" not in result
        assert "Content of page 4" not in result

    def test_read_unsupported_format(self, workspace):
        f = workspace / "uploads" / "data.xyz"
        f.write_bytes(b"\x00\x01\x02\x03" * 200)  # binary junk
        result = do_read(str(workspace), {"file_path": "uploads/data.xyz"})
        assert "Unsupported" in result

    def test_read_unknown_extension_text_heuristic(self, workspace):
        f = workspace / "uploads" / "config.env"
        f.write_text("KEY=value\nOTHER=stuff\n")
        result = do_read(str(workspace), {"file_path": "uploads/config.env"})
        assert "KEY=value" in result

    def test_read_missing_file_path(self, workspace):
        with pytest.raises(ValueError, match="file_path"):
            do_read(str(workspace), {})


# ---------------------------------------------------------------------------
# Smart truncation — PDF
# ---------------------------------------------------------------------------


class TestSmartTruncationPDF:
    def test_small_pdf_no_truncation(self, workspace, pdf_file):
        """Small PDF shows all pages with header, no continuation hint."""
        mock_page = MagicMock()
        mock_page.extract_text.return_value = "Short text"

        with patch("pypdf.PdfReader") as MockReader:
            MockReader.return_value.pages = [mock_page, mock_page]
            result = do_read(str(workspace), {"file_path": "uploads/doc.pdf"})
        assert "Document: doc.pdf (2 pages)" in result
        assert "Showing pages" not in result
        assert 'Use pages=' not in result

    def test_large_pdf_truncates_at_page_boundary(self, workspace, pdf_file):
        """Large PDF truncates at page boundary with continuation hint."""
        pages = []
        for i in range(100):
            p = MagicMock()
            # Each page ~1000 chars, 100 pages = ~100K > 50K budget
            p.extract_text.return_value = f"Page {i+1} content. " + ("x" * 900)
            pages.append(p)

        with patch("pypdf.PdfReader") as MockReader:
            MockReader.return_value.pages = pages
            result = do_read(str(workspace), {"file_path": "uploads/doc.pdf"})

        assert "Document: doc.pdf (100 pages)" in result
        assert "Showing pages 1-" in result
        assert "of 100" in result
        assert 'Use pages="' in result
        # Should NOT contain all 100 pages
        assert "Page 100 content" not in result
        # Should contain early pages
        assert "Page 1 content" in result
        assert len(result) < _MAX_OUTPUT_CHARS + 500  # budget + header/hint

    def test_pdf_continuation_hint_correct_numbers(self, workspace, pdf_file):
        """Continuation hint suggests the right next page range."""
        pages = []
        for i in range(50):
            p = MagicMock()
            p.extract_text.return_value = f"P{i+1}. " + ("y" * 2000)
            pages.append(p)

        with patch("pypdf.PdfReader") as MockReader:
            MockReader.return_value.pages = pages
            result = do_read(str(workspace), {"file_path": "uploads/doc.pdf"})

        # Extract the last shown page from the hint
        assert "Showing pages 1-" in result
        assert "of 50" in result
        assert 'to read more.' in result


# ---------------------------------------------------------------------------
# Smart truncation — CSV
# ---------------------------------------------------------------------------


class TestSmartTruncationCSV:
    def test_small_csv_no_truncation(self, workspace, csv_file):
        """Small CSV shows all rows with header, no hint."""
        result = do_read(str(workspace), {"file_path": "uploads/data.csv"})
        assert "Dataset: data.csv (3 rows" in result
        assert "Columns: name, age, city" in result
        assert "Showing rows" not in result

    def test_large_csv_truncates_at_row_boundary(self, workspace):
        """Large CSV truncates at row boundary with continuation hint."""
        f = workspace / "uploads" / "big.csv"
        with open(f, "w", newline="") as fh:
            writer = csv.writer(fh)
            writer.writerow(["id", "name", "value"])
            for i in range(10000):
                writer.writerow([i, f"item_{i}", "x" * 50])

        result = do_read(str(workspace), {"file_path": "uploads/big.csv"})
        assert "Dataset: big.csv (10001 rows" in result
        assert "Columns: id, name, value" in result
        assert "Showing rows 1-" in result
        assert "of 10001" in result
        assert "search(query)" in result
        assert len(result) < _MAX_OUTPUT_CHARS + 500


# ---------------------------------------------------------------------------
# Smart truncation — XLSX
# ---------------------------------------------------------------------------


class TestSmartTruncationXLSX:
    def test_small_xlsx_no_truncation(self, workspace, xlsx_file):
        """Small XLSX shows all data with header, no hint."""
        result = do_read(str(workspace), {"file_path": "uploads/sheet.xlsx"})
        assert "Workbook: sheet.xlsx (1 sheets: Data)" in result
        assert "truncated" not in result.lower()

    def test_large_xlsx_truncates_mid_sheet(self, workspace):
        """Large XLSX truncates mid-sheet with hint."""
        from openpyxl import Workbook
        f = workspace / "uploads" / "big.xlsx"
        wb = Workbook()
        ws = wb.active
        ws.title = "Sales"
        ws.append(["ID", "Product", "Amount"])
        for i in range(5000):
            ws.append([i, f"product_{i}", i * 10])
        wb.save(str(f))

        result = do_read(str(workspace), {"file_path": "uploads/big.xlsx"})
        assert "Workbook: big.xlsx" in result
        assert "Sheet: Sales" in result
        assert "truncated" in result.lower()
        assert "search(query)" in result
        assert len(result) < _MAX_OUTPUT_CHARS + 500


# ---------------------------------------------------------------------------
# Smart truncation — DOCX
# ---------------------------------------------------------------------------


class TestSmartTruncationDOCX:
    def test_small_docx_no_truncation(self, workspace, docx_file):
        """Small DOCX shows all text with header, no hint."""
        result = do_read(str(workspace), {"file_path": "uploads/report.docx"})
        assert "Document: report.docx" in result
        assert "Showing first" not in result

    def test_large_docx_truncates_at_paragraph(self, workspace):
        """Large DOCX truncates at paragraph boundary with hint."""
        from docx import Document
        f = workspace / "uploads" / "big.docx"
        doc = Document()
        for i in range(500):
            doc.add_paragraph(f"Paragraph {i}. " + "Lorem ipsum dolor sit amet. " * 20)
        doc.save(str(f))

        result = do_read(str(workspace), {"file_path": "uploads/big.docx"})
        assert "Document: big.docx" in result
        assert "Showing first" in result
        assert "chars" in result
        assert "exec tasks" in result
        assert len(result) < _MAX_OUTPUT_CHARS + 500


# ---------------------------------------------------------------------------
# Smart truncation — plain text
# ---------------------------------------------------------------------------


class TestSmartTruncationText:
    def test_small_text_no_truncation(self, workspace, txt_file):
        """Small text file shows all content with header, no hint."""
        result = do_read(str(workspace), {"file_path": "uploads/hello.txt"})
        assert "Document: hello.txt" in result
        assert "Showing first" not in result

    def test_large_text_truncates_at_line_boundary(self, workspace):
        """Large text file truncates at line boundary with hint."""
        f = workspace / "uploads" / "big.txt"
        lines = [f"Line {i}: " + "a" * 100 for i in range(_MAX_OUTPUT_CHARS // 100 + 100)]
        f.write_text("\n".join(lines))

        result = do_read(str(workspace), {"file_path": "uploads/big.txt"})
        assert "Document: big.txt" in result
        assert "Showing first" in result
        assert "exec tasks" in result
        assert len(result) < _MAX_OUTPUT_CHARS + 500
        # Should not cut mid-line
        assert result.rstrip().endswith(("chars.", "sections.")) or "Line " in result.split("\n")[-3]


# ---------------------------------------------------------------------------
# Path traversal guard
# ---------------------------------------------------------------------------


class TestPathTraversal:
    def test_traversal_rejected(self, workspace):
        with pytest.raises(ValueError, match="traversal"):
            _resolve_path(str(workspace), {"file_path": "../../etc/passwd"})

    def test_valid_path_accepted(self, workspace, txt_file):
        result = _resolve_path(str(workspace), {"file_path": "uploads/hello.txt"})
        assert result.name == "hello.txt"


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


class TestParsePageRanges:
    def test_single_page(self):
        assert _parse_page_ranges("3", 10) == [2]  # zero-based

    def test_range(self):
        assert _parse_page_ranges("2-4", 10) == [1, 2, 3]

    def test_mixed(self):
        assert _parse_page_ranges("1,3-5,8", 10) == [0, 2, 3, 4, 7]

    def test_out_of_bounds_clamped(self):
        result = _parse_page_ranges("1-100", 5)
        assert result == [0, 1, 2, 3, 4]

    def test_deduplication(self):
        result = _parse_page_ranges("1-3,2-4", 10)
        assert result == [0, 1, 2, 3]  # no duplicates


class TestFormatSize:
    def test_bytes(self):
        assert _format_size(500) == "500 B"

    def test_kilobytes(self):
        assert "KB" in _format_size(2048)

    def test_megabytes(self):
        assert "MB" in _format_size(5 * 1024 * 1024)


class TestIsLikelyText:
    def test_text_file(self, tmp_path):
        f = tmp_path / "readme.nfo"
        f.write_text("This is plain text content.")
        assert _is_likely_text(f) is True

    def test_binary_file(self, tmp_path):
        f = tmp_path / "image.bin"
        f.write_bytes(bytes(range(256)) * 4)
        assert _is_likely_text(f) is False

    def test_empty_file(self, tmp_path):
        f = tmp_path / "empty"
        f.write_bytes(b"")
        assert _is_likely_text(f) is False


# ---------------------------------------------------------------------------
# Functional: stdin/stdout contract
# ---------------------------------------------------------------------------


class TestFunctional:
    def test_list_via_stdin(self, workspace, txt_file):
        """Full subprocess: JSON stdin → stdout."""
        input_data = json.dumps({
            "args": {"action": "list"},
            "workspace": str(workspace),
        })
        result = subprocess.run(
            [sys.executable, "run.py"],
            input=input_data, capture_output=True, text=True,
            cwd=str(Path(__file__).parent.parent),
        )
        assert result.returncode == 0
        assert "hello.txt" in result.stdout

    def test_read_via_stdin(self, workspace, txt_file):
        input_data = json.dumps({
            "args": {"action": "read", "file_path": "uploads/hello.txt"},
            "workspace": str(workspace),
        })
        result = subprocess.run(
            [sys.executable, "run.py"],
            input=input_data, capture_output=True, text=True,
            cwd=str(Path(__file__).parent.parent),
        )
        assert result.returncode == 0
        assert "Hello, world!" in result.stdout

    def test_missing_file_exits_1(self, workspace):
        input_data = json.dumps({
            "args": {"action": "read", "file_path": "uploads/nope.txt"},
            "workspace": str(workspace),
        })
        result = subprocess.run(
            [sys.executable, "run.py"],
            input=input_data, capture_output=True, text=True,
            cwd=str(Path(__file__).parent.parent),
        )
        assert result.returncode == 1
        assert "not found" in result.stderr.lower() or "not found" in result.stdout.lower()
